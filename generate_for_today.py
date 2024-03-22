"""
By Michał Matuszyk
on 20/03/2024
"""
import random

from docx import Document
import datetime
import os


list_of_quotes = [
 """"It takes courage to grow up and become who you really are." — E.E. Cummings""",
    """"There is nothing impossible to they who will try."— Alexander the Great""",
    """"Live a life you will remember" - Avicii"""
]




def replace_text_in_docx(docx_file, replacement_dict):
    doc = Document(docx_file)

    for paragraph in doc.paragraphs:
        for key, value in replacement_dict.items():
            if key in paragraph.text:
                paragraph.text = paragraph.text.replace(key, value)

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    for key, value in replacement_dict.items():
                        if paragraph.text == key:
                            paragraph.text = paragraph.text.replace(key, value)
    return doc

def populate_dates_table():
    number_of_weeks = 5
    today = datetime.date.today()
    days_to_last_sunday = today.weekday() - 6  # 6 represents Sunday
    last_sunday = today - datetime.timedelta(days=days_to_last_sunday + 7)
    start_date = last_sunday# - datetime.timedelta(days=time_delta - 1)
    current_date = start_date
    populate_dict = {}
    for i in range(0, number_of_weeks * 7):
        populate_dict[chr((i//7 + 1 + 64)) + chr((i%7+1) + 64)] = str(current_date.day)
        current_date += datetime.timedelta(days=1)
    return populate_dict

def get_quote():

    if "quotes.txt" in os.listdir("."):
        print("Found quotes file - using it")
        with open("quotes.txt", 'r') as f:
            quotes = f.readlines()
            return quotes[random.randrange(len(quotes))]

    return list_of_quotes[random.randrange(len(list_of_quotes))]

def create_folder_and_file():
    tomorrow = datetime.date.today() + datetime.timedelta(days=1)
    folder_name = tomorrow.strftime("%Y_%m_%d")
    os.makedirs(folder_name, exist_ok=True)
    file_path = os.path.join(folder_name, "things.txt")
    with open("things_template.txt", "r") as template_file:
        template_content = template_file.read()

    with open(file_path, "w") as new_file:
        new_file.write(template_content)
    print(f"Folder '{folder_name}' and file 'things.txt' created successfully.")
    return folder_name



def create_todo_dict(todos):
    todo_count = 19
    result_dict = {}
    for i in range(todo_count):
        if len(todos) > i:
            result_dict['Todo' + str(i+1)] = todos[i]
        else:
            result_dict['Todo' + str(i+1)] = ""
    return result_dict

def create_habit_dict(habits):
    habit_count = 7
    result_dict = {}
    for i in range(habit_count):
        if len(habits) > i:
            result_dict['Habit_' + str(i + 1)] = habits[i]
        else:
            result_dict['Habit_' + str(i + 1)] = ""
    return result_dict

def parse_file(file_path):
    with open(file_path, 'r', encoding='utf-8') as f:
        lines = f.readlines()
    Quote = get_quote()
    todo_items = []
    habit_items = []
    for line in lines:
        if line[0] == "%":
            continue
        elif line[0] == "T":
            todo_items.append(line[2:])
        elif line[0] == "H":
            habit_items.append(line[2:])
        elif line[0] == "Q":
            Quote = line[2:]
    return Quote, todo_items, habit_items

def save_as_pdf(docx, pdf_filename):
    docx.save(pdf_filename)



def main():
    created_folder_path = create_folder_and_file()
    tomorrow = datetime.date.today() + datetime.timedelta(days=1)
    formatted_date = tomorrow.strftime("%Y_%m_%d")
    input("Press enter when ready")
    Quote, todo_items, habit_items = parse_file(created_folder_path + "/things.txt")
    # Replacement dictionary
    replacement_dict = {
        "DD/MM/YYYY": tomorrow.strftime("%d/%m/%Y"),
        "Day_of_week": tomorrow.strftime('%A'),
        "QUOTE": Quote,
        "MONTH_PLACEHOLDER": "March"
    }
    days = populate_dates_table()
    replacement_dict = {**replacement_dict, **days}
    todos_dict = create_todo_dict(todo_items)
    replacement_dict = {**replacement_dict, **todos_dict}
    habits_dict = create_habit_dict(habit_items)
    replacement_dict = {**replacement_dict, **habits_dict}
    docx_file = "day_template.docx"
    doc = replace_text_in_docx(docx_file, replacement_dict)
    pdf_filename = f"{created_folder_path}/{formatted_date}.pdf"
    save_as_pdf(doc, pdf_filename)
    print(f"PDF saved as {pdf_filename}")

if __name__ == "__main__":
    main()
